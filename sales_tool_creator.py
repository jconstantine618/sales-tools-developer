import streamlit as st
from typing import Dict, List, Tuple
from openai import OpenAI
from docx import Document
from docx.shared import Pt
import requests, bs4
from urllib.parse import urljoin, urlparse

# ────────────────────────────────────────────────────────────────────────────────
# PAGE CONFIG
# ────────────────────────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="B2B Sales Playbook Builder", page_icon="📘", layout="wide"
)

# ────────────────────────────────────────────────────────────────────────────────
# OPENAI CLIENT (cached per session)
# ────────────────────────────────────────────────────────────────────────────────

@st.cache_resource(show_spinner=False)
def get_openai() -> OpenAI:
    return OpenAI(api_key=st.secrets["openai_api_key"])

client: OpenAI = get_openai()

# ────────────────────────────────────────────────────────────────────────────────
# CONSTANTS
# ────────────────────────────────────────────────────────────────────────────────

SECTION_TITLES: List[str] = [
    "Company Overview",
    "Value Propositions",
    "Customer Benefits",
    "Target Audience",
    "Needs Assessment Questions",
    "Demo Customer Personas",
    "Closing Questions",
    "Lead Generation Channels & Next Steps",
]

MAX_SITE_PAGES = 10  # safety‑limit for crawler
MAX_SITE_CHARS = 5000  # send a sane chunk to GPT

# ────────────────────────────────────────────────────────────────────────────────
# WEBSITE SCRAPER
# ────────────────────────────────────────────────────────────────────────────────

@st.cache_data(show_spinner="🌐 Gathering website copy …")
def scrape_public_site(root_url: str, max_pages: int = MAX_SITE_PAGES) -> str:
    """Grab visible text from up to *max_pages* internal URLs (DFS crawl)."""

    domain = urlparse(root_url).netloc
    seen, to_visit = set(), [root_url]
    texts: List[str] = []

    while to_visit and len(seen) < max_pages:
        url = to_visit.pop(0)
        if url in seen:
            continue
        seen.add(url)
        try:
            r = requests.get(url, timeout=6)
            if not r.ok or "text/html" not in r.headers.get("Content-Type", ""):
                continue
        except Exception:
            continue

        soup = bs4.BeautifulSoup(r.text, "html.parser")
        for s in soup(["script", "style", "noscript"]):
            s.extract()
        page_text = " ".join(t.strip() for t in soup.stripped_strings)
        texts.append(page_text)

        for a in soup.find_all("a", href=True):
            link = urljoin(url, a["href"])
            if (
                urlparse(link).netloc == domain
                and link not in seen
                and link not in to_visit
            ):
                to_visit.append(link)

    return " \n".join(texts)[:MAX_SITE_CHARS]

# ────────────────────────────────────────────────────────────────────────────────
# GPT‑4 SECTION GENERATION
# ────────────────────────────────────────────────────────────────────────────────

def _persona_bullets(personas: List[Dict]) -> str:
    if not personas:
        return "N/A"
    return "\n".join([
        f"- {p['industry']} {p['persona']}: {p['relation']}" for p in personas
    ])


def generate_section_content(section: str, info: Dict, personas: List[Dict]) -> str:
    base = f"""
Company Name: {info['company_name']}
Products/Services: {info['products_services']}
Target Audience: {info['target_audience']}
Top Problems: {info['top_problems']}
Unique Value Proposition: {info['value_prop']}
Website Copy Excerpt: {info['website_text'][:1500]}
Personas:\n{_persona_bullets(personas)}
Tone: {info['tone']}
"""
    prompt = f"""{base}

Write the **{section}** section of a B2B Sales Playbook. Adopt a professional yet conversational tone influenced by Dale Carnegie, Challenger, and Sandler methodologies. Use clear sub‑headers and bullet points where useful."""

    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": "You are a B2B sales strategist writing sales playbooks based on company profiles.",
            },
            {"role": "user", "content": prompt},
        ],
        temperature=0.7,
    )
    return resp.choices[0].message.content.strip()


def generate_all_sections(info: Dict, personas: List[Dict]) -> Dict[str, str]:
    """Generate every playbook section on‑demand. No cache so we only run when user clicks."""
    return {sec: generate_section_content(sec, info, personas) for sec in SECTION_TITLES}

# ────────────────────────────────────────────────────────────────────────────────
# WORD EXPORT
# ────────────────────────────────────────────────────────────────────────────────

def build_word_doc(company_name: str, section_texts: Dict[str, str]) -> Document:
    doc = Document()
    doc.add_heading(f"{company_name} B2B Sales Playbook", 0)
    for title, body in section_texts.items():
        doc.add_heading(title, level=1)
        for para in body.split("\n"):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.style.font.size = Pt(11)
    return doc

# ────────────────────────────────────────────────────────────────────────────────
# SIDEBAR INPUTS
# ────────────────────────────────────────────────────────────────────────────────

def sidebar_inputs() -> Tuple[Dict, List[Dict], bool]:
    """Collect user inputs and return info, personas, and whether the generate button was pressed."""
    with st.sidebar:
        st.header("Company Info")
        st.caption("Fill out the fields below. **Nothing is generated until you press the 'Generate / Update Playbook' button at the bottom.**")

        company_name = st.text_input("Company Name")
        products_services = st.text_area("Products / Services")
        target_audience = st.text_input("Target Audience")
        top_problems = st.text_area("Top Problems")
        value_prop = st.text_area("Unique Value Proposition")
        website_url = st.text_input("Company Website URL (https://…)")
        tone = st.selectbox(
            "Writing Tone",
            ["Professional", "Friendly", "Conversational", "Challenger"],
            index=0,
        )

        # Crawl website when the user explicitly requests it (optional)
        if website_url and st.button("🔍 Fetch website copy", key="fetch_site"):
            with st.spinner("Scraping website…"):
                st.session_state.website_text = scrape_public_site(website_url)

        st.divider()
        st.header("Prospect Types (max 5)")

        if "num_personas" not in st.session_state:
            st.session_state.num_personas = 1

        if st.button("➕ Add another prospect type", disabled=st.session_state.num_personas >= 5):
            st.session_state.num_personas += 1

        personas: List[Dict] = []
        for i in range(st.session_state.num_personas):
            with st.expander(f"Prospect {i+1}"):
                industry = st.text_input("Company / Industry", key=f"pers_{i}_industry")
                role = st.text_input("Role / Title", key=f"pers_{i}_role")
                relation = st.text_area(
                    "Why this role cares about your service", key=f"pers_{i}_relation", height=80
                )
                if industry or role or relation:
                    personas.append(
                        {"industry": industry, "persona": role, "relation": relation}
                    )

        st.markdown("---")
        generate_clicked = st.button("🚀 Generate / Update Playbook")
        if generate_clicked:
            st.success("Generating or updating your playbook… you can keep working; results will appear in the main panel.")

    info = {
        "company_name": company_name,
        "products_services": products_services,
        "target_audience": target_audience,
        "top_problems": top_problems,
        "value_prop": value_prop,
        "website_text": st.session_state.get("website_text", ""),
        "tone": tone,
    }
    return info, personas, generate_clicked

# ────────────────────────────────────────────────────────────────────────────────
# RENDERER
# ────────────────────────────────────────────────────────────────────────────────

def render_playbook_builder():
    st.markdown("## 📘 Build Your Custom B2B Sales Playbook")

    info, personas, generate_clicked = sidebar_inputs()

    # Only (re)generate when the user explicitly presses the button
    if generate_clicked:
        st.session_state.playbook_sections = generate_all_sections(info, personas)
        st.session_state.playbook_company = info["company_name"] or "Company"

    if "playbook_sections" in st.session_state:
        company_header = st.session_state.get("playbook_company", "Company")
        st.markdown(f"### ✨ {company_header} Playbook Preview")

        # Editable text areas for each section
        for section in SECTION_TITLES:
            content = st.session_state.playbook_sections.get(section, "")
            new_content = st.text_area(section, value=content, height=200)
            st.session_state.playbook_sections[section] = new_content

        # Export button
        if st.button("💾 Export to Word (.docx)"):
            doc = build_word_doc(company_header, st.session_state.playbook_sections)
            file_name = f"{company_header}_Sales_Playbook.docx"
            doc.save(file_name)
            with open(file_name, "rb") as f:
                st.download_button(
                    label="Download Playbook", data=f, file_name=file_name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    else:
        st.info("Fill out information in the sidebar, then click **Generate / Update Playbook** to create your playbook. You can tweak any field later and press the button again to refresh the content – the playbook will not regenerate automatically while you type.")

# ────────────────────────────────────────────────────────────────────────────────
# MAIN
# ────────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    render_playbook_builder()
